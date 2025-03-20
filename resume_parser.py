import os
import docx
from PyPDF2 import PdfReader

def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, "rb") as file:
        reader = PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        return "Unsupported file format"

# ✅ **Test the function**
if __name__ == "__main__":
    resume_text = extract_text("Saikumar_Mehtre_MisExecutive.pdf")
    print(resume_text)

import PyPDF2
import docx
import io  # Add this to handle in-memory files

def extract_text_from_pdf(pdf_file):
    """Extract text from an uploaded PDF file."""
    reader = PyPDF2.PdfReader(pdf_file)
    text = "".join([page.extract_text() for page in reader.pages if page.extract_text()])
    return text

def extract_text_from_docx(docx_file):
    """Extract text from an uploaded DOCX file."""
    doc = docx.Document(docx_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text


def extract_skills(text):
    skills = ["Python", "SQL", "Machine Learning", "Tableau", "Excel"]
    extracted_skills = [skill for skill in skills if skill.lower() in text.lower()]
    return extracted_skills